import streamlit as st
from docx import Document
from docx.shared import Inches
import fitz
import re
import tempfile

st.title("📄 Gerador de Provas de Física")

tema = st.text_input("Tema (ex: velocidade média)")
quantidade = st.slider("Quantidade de questões", 1, 10, 5)

pdfs = st.file_uploader("Envie PDFs com questões", accept_multiple_files=True)
modelo = st.file_uploader("Modelo Word com {{QUESTOES}}")

def separar_questoes(texto):
    partes = re.split(r"\n?\s*(\d+)[\.\)]\s+", texto)
    questoes = []
    for i in range(1, len(partes), 2):
        q = partes[i+1].strip()
        if len(q) > 40:
            questoes.append(q)
    return questoes

def ler_pdf(file):
    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(file.read())
        caminho = tmp.name

    doc = fitz.open(caminho)
    texto = ""
    for p in doc:
        texto += p.get_text()
    return texto

if st.button("Gerar prova"):
    if not pdfs or not modelo:
        st.warning("Envie os PDFs e o modelo Word")
    else:
        todas_questoes = []

        for pdf in pdfs:
            texto = ler_pdf(pdf)
            qs = separar_questoes(texto)

            for q in qs:
                if tema.lower() in q.lower():
                    todas_questoes.append(q)

        selecionadas = todas_questoes[:quantidade]

        doc = Document(modelo)

        for p in doc.paragraphs:
            if "{{QUESTOES}}" in p.text:
                p.text = ""
                for i, q in enumerate(selecionadas, 1):
                    doc.add_paragraph(f"{i}. {q}")
                    doc.add_paragraph("")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            st.download_button("📥 Baixar prova", open(tmp.name, "rb"), file_name="prova.docx")
